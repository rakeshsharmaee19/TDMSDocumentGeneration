import matplotlib.pyplot as plt
from matplotlib.ticker import (MultipleLocator)
from docx.shared import Cm, Pt
from .dataFromJSON import *


### <<<< -------------------------- Generate Graph Using Graph Data ------------------------- >>>>
def generateGraph(data, heading):
    x = (
    0, 0.25, 0.50, 0.75, 1, 1.25, 1.50, 1.75, 2, 2.25, 2.50, 2.75, 3, 3.25, 3.50, 3.75, 4, 4.25, 4.50, 4.75, 5, 5.25,
    5.50, 5.75, 6, 6.25, 6.50, 6.75, 7, 7.25, 7.50, 7.75, 8, 8.25, 8.50, 8.75, 9, 9.25, 9.50, 9.75, 10)
    fig, ax = plt.subplots(layout='constrained')
    ax.plot(x, data[0], 'b-o', label='Trend Line')
    ax.plot(x, data[1], 'gD', label='Measured Resistances')
    ax.xaxis.set_major_locator(MultipleLocator(1))
    ax.xaxis.set_minor_locator(MultipleLocator(0.25))
    ax.margins(x=0.001)
    ax.grid(True)
    ax.set_xlabel(' Time (Min) ', fontsize=20)
    ax.set_ylabel(' Hot Resistance Curve (Ohms) ', fontsize=16)
    ax.legend()
    ax.set_title(heading)
    fig.savefig('picture/f.png')
    # fig.savefig('C:\\Users\\Rakesh\\PycharmProjects\\TDMSDocumentGeneration\\picture\\f.png')


### <<<< ------------------------------ Add Image in Table Cell ------------------------------ >>>>
def addGraph(table, rowNumber=0):
    p = table.rows[rowNumber].cells[0].add_paragraph()
    r = p.add_run()
    r.add_picture('picture/f.png', width=Cm(9), height=Cm(6))


### <<<< ----------------------------- Adding Heading For The Graph ------------------------------ >>>>

# def graph_heading(data, windingKey):
#     graphHeading = data['DATA'][windingKey]["hotResistance"]["title"]
#     if graphHeading == None:
#         graphHeading = windingKey.upper()
#     else:
#         graphHeading = graphHeading + '  ' + windingKey.upper()
#     return graphHeading

def graph_heading(data, windingKey):
    graphHeading = data['DATA'][windingKey]['transformerId'] + ' ' + data['DATA'][windingKey]['hrMVA'] + ' ' + '(' + data['DATA'][windingKey]['basicInfo']['cooling'] + ')' + ' ' + windingKey
    return graphHeading
