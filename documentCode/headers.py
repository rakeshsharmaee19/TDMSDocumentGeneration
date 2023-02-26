from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT

from .styleTable import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

### <<<< ------------------------ Generate Header Data --------------------------- >>>>
def headerData(data):
    data1 = []
    headerKeys = [i for i in data.keys()]
    if len(headerKeys) == 9:
        for i in range(len(headerKeys)):
            if i == 0 or i == 1 or i == 8:
                data1.append(data[headerKeys[i]])
            elif i in (2, 3, 7):
                keyData = headerKeys[i].upper() + ':  '
                data1.append(keyData)
                data1.append(data[headerKeys[i]])
            elif i == 4:
                data1.append("CUSTOMER PURCHASE ORDER NUMBER:   ")
                data1.append(data[headerKeys[i]])
            elif i == 5:
                data1.append("PROLEC GE WAUKESHA, INC.ORDER NO:  ")
                data1.append(data[headerKeys[i]])
            else:
                data1.append("SERIAL NUMBER: ")
                data1.append(data[headerKeys[i]])
        return data1
    else:
        return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": "list index out of range",
                    "errorMessage": "Header Data is not proper, key is missing"}

### <<<< ------------------------------ Style Header Table ---------------------------------- >>>>>
def headerTableStyle(table, cell1 = 12, cell2 = 6):
    table.style = 'TableGrid'
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in range(len(table.rows)):
        rowHeight(table.rows[row], 0.6)
        if row in (0,2,4):
            a = table.cell(row, 0)
            b = table.cell(row, 1)
            a.merge(b)
            table.rows[row].cells[0].width = Cm(cell1)
            table.rows[row].cells[2].width = Cm(cell2)
            if row == 0:
                rowHeight(table.rows[row], 1.8)
        elif row in (1,3):
            a = table.cell(row, 0)
            b = table.cell(row, 1)
            c = table.cell(row, 2)
            d = a.merge(b)
            d.merge(c)
        else:
            a = table.cell(row, 1)
            b = table.cell(row, 2)
            a.merge(b)
            table.rows[row].cells[0].width = Cm(6)
            table.rows[row].cells[2].width = Cm(12)
            rowHeight(table.rows[row], 1)

### <<<< ------------------------------ Fill Value in Header Table ----------------------------- >>>>
def fillHeaderValue(table, value):
    count = 0
    for i in range(len(table.rows)):
        if i == 0:
            makeBoldItalic(table, rowNumber=i, cellNumber=0, bold=True, italic=True, fontSize=14, text=value[count])
            table.cell(i, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            p = table.rows[i].cells[2].add_paragraph()
            r = p.add_run()
            r.add_picture('picture/logo.png', width=Cm(4), height=Cm(1.26))
            table.cell(i, 2).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            count+=1
        elif i == 1:
            makeBoldItalic(table, rowNumber=i, cellNumber=0, bold=True, italic=False, fontSize=14, text=value[count])
            count+=1
        elif i in (2,4,5):
            try:
                text =value[count]+value[count+1]
            except Exception as e:
                return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed",
                               "error": str(e),
                               "errorMessage": "Header Data is not proper, Key Value is Null"}
            makeBoldItalic(table, rowNumber=i, cellNumber=0, bold=True, italic=False, fontSize=11, text=text)
            count+=2
            if i == 5:
                try:
                    # newText = 'value[count][0:2]+value[count][3:]'
                    replaceValue = '\u00B0' + 'C '
                    # if 'C' in value[count] or '65' in value[count]:
                    #     replaceText = value[count].replace('65C', replaceValue)
                    # elif '65' in value[count]:
                    #     replaceText = value[count].replace('65', replaceValue)
                    if "$C" in value[count]:
                        replaceText = value[count].replace('$C', replaceValue)
                    else:
                        replaceText= value[count]
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed",
                            "error": str(e),
                            "errorMessage": "Header Data is not proper, Key Value is Null"}
                makeBoldItalic(table, rowNumber=i, cellNumber=2, bold=True, italic=False, fontSize=11,
                               text=replaceText)
                count += 1
                table.cell(i, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            else:
                try:
                    text = value[count] + value[count + 1]
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed",
                            "error": str(e),
                            "errorMessage": "Header Data is not proper, Key Value is Null"}
                makeBoldItalic(table, rowNumber=i, cellNumber=2, bold=True, italic=False, fontSize=11,
                               text=text)
                count += 2
        else:
            try:
                text = value[count] + value[count + 1]
            except Exception as e:
                return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed",
                        "error": str(e),
                        "errorMessage": "Header Data is not proper, Key Value is Null"}
            makeBoldItalic(table, rowNumber=i, cellNumber=0, bold=True, italic=False, fontSize=11, text=text)
            count+=2

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def createHeaderFooter(document, orientation=False, width = 18, cell1=12, cell2=6):
    if orientation:
        section = change_orientation(document)
    else:
        section = document.sections[0]
    header = section.header
    headerTable = header.add_table(rows=6, cols=3, width=Cm(width))
    headerTableStyle(headerTable, cell1 = cell1, cell2=cell2)
    header.add_paragraph('')
    add_page_number(document.sections[0].footer.paragraphs[0].add_run())
    document.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    footer = section.footer
    footer_para = footer.paragraphs[0].add_run()
    footer_para.text = "                            © 2023 Prolec GE Waukesha, Inc.\n\tAll Rights Reserved"
    return headerTable