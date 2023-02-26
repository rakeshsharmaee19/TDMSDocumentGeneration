from docx.shared import Cm, Pt, Inches, Mm
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START

### <<<< ------------------------------ Change Table Row height ------------------------------->>>>
def rowHeight(table, value):
    table.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    table.height = Cm(value)

def set_col_widths(table, columnWidth=[]):
    cellWidth = columnWidth
    for row in table.rows:
        rowHeight(row, 0.5)
        for idx, width in enumerate(cellWidth):
            row.cells[idx].width = width

def makeBoldItalic(table, rowNumber=0, cellNumber=0, bold=True, italic=True, fontSize=12, text="textValue"):
    p = table.rows[rowNumber].cells[cellNumber].paragraphs[0]
    runner = p.add_run(text)
    font = runner.font
    font.name = 'Times New Roman'
    font.size = Pt(fontSize)
    runner.bold = bold
    runner.italic = italic

def makeBoldPara(paragraph, bold=True, italic=False, fontSize=14, text='', alignment=0):
    p = paragraph
    p.alignment = alignment
    runner = p.add_run(text)
    font = runner.font
    font.name = 'Times New Roman'
    font.size = Pt(fontSize)
    runner.bold = bold
    runner.italic = italic

def hideBorder(table):  ## hide border
    tbl = table._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        tcBorders = OxmlElement("w:tcBorders")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "nil")

        left = OxmlElement("w:left")
        left.set(qn("w:val"), "nil")

        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "nil")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), "auto")

        right = OxmlElement("w:right")
        right.set(qn("w:val"), "nil")

        tcBorders.append(top)
        tcBorders.append(left)
        tcBorders.append(bottom)
        tcBorders.append(right)
        tcPr.append(tcBorders)

def styeleOverAllTable(table):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    #     #Merging the columns :
    table.cell(0, 0).merge(table.cell(0, 5))
    table.cell(1, 0).merge(table.cell(2, 0))
    table.cell(1, 1).merge(table.cell(2, 1))
    table.cell(1, 2).merge(table.cell(2, 2))
    table.cell(1, 3).merge(table.cell(2, 3))
    table.cell(1, 4).merge(table.cell(1, 5))

def styleBushingTable(table, merge_list = [], num_of_columns=0):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.cell(0, 0).merge(table.cell(0, num_of_columns - 1))
    k = 3
    for j in merge_list:
        for i in range(k, num_of_columns - 1, j):
            table.cell(1, i + 1).merge(table.cell(1, i + j))
            break
        k = k + j
    for i in range(0, 4):
        table.cell(1, i).merge(table.cell(2, i))

def styleBelowTable(table):
    for row in range(len(table.rows)):
        rowHeight(table.rows[row], 0.5)

def styleExcitationTestTable(table, numberOfColumns):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cellWidth = [Cm(2), Cm(4.5), Cm(4.5), Cm(2), Cm(2.9), Cm(1.8)]
    set_col_widths(table, columnWidth=cellWidth)
    table.cell(0, 0).merge(table.cell(0, numberOfColumns - 1))

def styleExcitationTable(table, numberOfColumns, mergeList=[]):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.cell(0, 0).merge(table.cell(0, numberOfColumns - 1))
    if numberOfColumns == 5:
        cellWidth =[Cm(2), Cm(2.5), Cm(2.5), Cm(2), Cm(2.5)]
        set_col_widths(table, columnWidth=cellWidth)
    else:
        cellWidth = [Cm(1.8), Cm(2.3), Cm(2), Cm(1.8), Cm(1.8), Cm(1.8), Cm(2.3), Cm(2), Cm(1.8), Cm(1.8)]
        set_col_widths(table, columnWidth=cellWidth)
    k = 0
    for j in mergeList:
        for i in range(k, numberOfColumns - 1, j):
            table.cell(1, i).merge(table.cell(1, i + j - 1))
            break
        k = k + j

# def styleRatioTest(table, numberOfColumns):
#     table.autofit = False
#     table.allow_autofit = False
#     table.alignment = WD_TABLE_ALIGNMENT.CENTER
#     table.cell(0, 0).merge(table.cell(0, numberOfColumns - 1))
#     for i in range(len(table.rows)):
#         if i == 1:
#             rowHeight(table.rows[i], 1)
#         else:
#             rowHeight(table.rows[i], 0.5)

def styleRatioTest(table, numberOfColumns=9, key=True):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    if numberOfColumns == 9:
        cellWidth = [Cm(1.6), Cm(1.6), Cm(2.8), Cm(2.4), Cm(2.4), Cm(2.4), Cm(1.5), Cm(1.5), Cm(1.5)]
        set_col_widths(table, columnWidth=cellWidth)

    if key:
        table.cell(0, 0).merge(table.cell(0, numberOfColumns - 1))
        table.cell(1, 0).merge(table.cell(1, numberOfColumns - 1))
        for i in range(len(table.rows)):
            if i == 1:
                rowHeight(table.rows[i], 1)
            else:
                rowHeight(table.rows[i], 0.5)
        table.cell(2, 0).merge(table.cell(2, numberOfColumns - 1))
        table.cell(3, 0).merge(table.cell(3, 2))
        table.cell(3, 3).merge(table.cell(3, 5))
        table.cell(3, 6).merge(table.cell(3, 8))
        table.cell(4, 0).merge(table.cell(4, 1))
        table.cell(4, 3).merge(table.cell(4, 5))
        table.cell(4, 6).merge(table.cell(4, 8))
        table.cell(4, 2).merge(table.cell(5, 2))

    elif not key:
        table.cell(0, 0).merge(table.cell(0, numberOfColumns - 1))
        table.cell(1, 0).merge(table.cell(1, 2))
        table.cell(1, 3).merge(table.cell(1, 5))
        table.cell(1, 6).merge(table.cell(1, 8))
        table.cell(2, 0).merge(table.cell(2, 1))
        table.cell(2, 3).merge(table.cell(2, 5))
        table.cell(2, 6).merge(table.cell(2, 8))
        table.cell(2, 2).merge(table.cell(3, 2))

def set_vertical_cell_direction(cell: _Cell, direction: str):
    # direction: tbRl -- top to bottom, btLr -- bottom to top
    assert direction in ("tbRl", "btLr")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), direction)  # btLr tbRl
    tcPr.append(textDirection)

def styleDgaTest(table, mergeVal=1, widthData=[]):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    dataPass = [Cm(i) for i in widthData]
    # for i in widthData:
    #     dataPass.app
    set_col_widths(table, columnWidth=dataPass)
    for i in range(len(table.rows)):
        if i == 0:
            rowHeight(table.rows[i], 4)
        else:
            rowHeight(table.rows[i], 0.5*mergeVal)

def change_orientation(document):
    current_section = document.sections[-1]
    # new_width, new_height = current_section.page_height, current_section.page_width
    # new_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    current_section.orientation = WD_ORIENTATION.LANDSCAPE
    current_section.page_width = Mm(297)
    current_section.page_height = Mm(210)
    return current_section

def styleOuterTable(table):   ## Heatrun data container table.
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    table.rows[0].cells[0].width = Cm(10)
    table.rows[0].cells[1].width = Cm(10)
    hideBorder(table)

def styleTableTop(table):   # style heatrun top table
    table.style = 'TableGrid'
    table.autofit = False
    table.allow_autofit = False
    numberOfColumns = len(table.rows[0].cells)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in range(len(table.rows)):
        rowHeight(table.rows[row], 0.6)
        if row in (0, 1, 2, 9):
            for col in range(len(table.rows[row].cells)):
                if col == 0:
                    table.rows[row].cells[col].width = Cm(8)
                else:
                    table.rows[row].cells[col].width = Cm(2.1)
            table.cell(row, 1).merge(table.cell(row, numberOfColumns - 1))
        else:
            for col in range(len(table.rows[row].cells)):
                if col == 0:
                    table.rows[row].cells[col].width = Cm(7)
                else:
                    table.rows[row].cells[col].width = Cm(2.3)

def firstTestColumnWidth(numberOfColumns):
    lis = [6, 4, 4, 4, 4]
    su = 0

    for i in range(numberOfColumns, len(lis)):
        su = su + lis[i]

    new_widths = lis[0:numberOfColumns]
    new_widths = [round(i + (su / numberOfColumns), 2) for i in new_widths]
    dataPass = [Cm(i) for i in new_widths]
    return dataPass

def styleTextTable(document, table, data=[]):
    section = document.sections[0]
    section.left_margin = Cm(1.5)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_col_widths(table, columnWidth=[Cm(18)])
    for row in range(len(table.rows)):
        rowHeight(table.rows[row], 0.5)
        if row !=0:
            rowHeight(table.rows[row], 0.5*data[row-1])


def styleEfficienceTable(table, numberOfRows=6, numberOfColumns=6, merge_list=[2,2]):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Merging the columns :
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(1, 1))

    k = 2
    for j in merge_list:
        for i in range(k, numberOfColumns - 1, j):
            table.cell(0, i).merge(table.cell(0, i + j - 1))
            break
        k = k + j
    table.cell(numberOfRows - 1, 0).merge(table.cell(numberOfRows - 1, 1))
    for i in range(len(table.rows)):
        if i == (len(table.rows)-1):
            rowHeight(table.rows[i], 0.5)
        elif i>1:
            rowHeight(table.rows[i], 0.5*2)
        elif i == 0:
            rowHeight(table.rows[i], 0.5*3)
        else:
            rowHeight(table.rows[i], 0.5)

def styleControlTable(table):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT


    set_col_widths(table, columnWidth=[Cm(6), Cm(3), Cm(3)])

    table.cell(1, 0).merge(table.cell(2, 0))
    table.cell(0, 0).merge(table.cell(0, len(table.rows[0].cells)-1))
    count = 1
    while count<len(table.rows[0].cells):
        table.cell(1, count).merge(table.cell(1, count+1))
        count+=2
    for row in table.rows:
        rowHeight(row, 0.5)

def styleEfficenceOuterTable(table, row = 6):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    r = row+2
    set_col_widths(table, columnWidth=[Cm(9), Cm(9)])
    rowHeight(table.rows[0], 0.5*r)

def styleEfficenceTable(table):
    table.style = 'TableGrid'
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    if len(table.rows[2].cells)==2:
        set_col_widths(table, columnWidth=[Cm(3.5), Cm(3.5)])
    elif len(table.rows[2].cells)==3:
        set_col_widths(table, columnWidth=[Cm(2.5), Cm(2.5), Cm(2.5)])
    table.cell(0, 0).merge(table.cell(0, len(table.rows[0].cells)-1))
    for row in table.rows:
        rowHeight(row, 0.5)

def stylePDTest(table, mergeRow=[]):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    width = [1.3, 1.0, 1.8, 2.3, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0]
    numberOfColumns = len(table.rows[0].cells)
    leftSum = 0
    if numberOfColumns<len(width):
        leftSum = sum(width)-sum(width[:numberOfColumns])
    newWidth = []
    for i in range(numberOfColumns):
        newWidth.append(Cm(width[i]+round(leftSum/numberOfColumns,3)))
    set_col_widths(table, columnWidth=newWidth)
    mergeCount = 4
    table.cell(0, 0).merge(table.cell(0, len(table.rows[0].cells) - 1))
    table.cell(1, 0).merge(table.cell(1, len(table.rows[0].cells) - 1))
    for i in mergeRow:
        table.cell(2, mergeCount).merge(table.cell(2, mergeCount+i-1))
        mergeCount+=i
    table.cell(2, 0).merge(table.cell(3, 0))
    table.cell(2, 1).merge(table.cell(3, 1))
    table.cell(2, 2).merge(table.cell(3, 2))
    table.cell(2, 3).merge(table.cell(3, 3))




def styleImpulseTestTable(table, numberOfColumns):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_col_widths(table, columnWidth=[Cm(5), Cm(3), Cm(3), Cm(3), Cm(3)])
    table.cell(0, 0).merge(table.cell(0, numberOfColumns - 1))


def stylePotencialTestTable(table, numberOfColumns):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_col_widths(table, columnWidth=[Cm(3), Cm(2.5), Cm(2.5), Cm(4), Cm(2.5)])
    table.cell(0, 0).merge(table.cell(0, numberOfColumns - 1))


def styleParallelInducedVoltageTestTable(table, numberOfColumns):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_col_widths(table, columnWidth=[Cm(6), Cm(4), Cm(4), ])
    table.cell(0, 0).merge(table.cell(0, numberOfColumns - 1))


def styleSeriesInducedVoltageTestTable(table, numberOfColumns):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_col_widths(table, columnWidth=[Cm(6), Cm(4), Cm(4), ])
    table.cell(0, 0).merge(table.cell(0, numberOfColumns - 1))


def styleInsulationResistanceTable(table, numberOfColumns, numOfHeaders):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    if numOfHeaders == 2:
        table.cell(0, 0).merge(table.cell(0, numberOfColumns-1))
    else:
        table.cell(0, 0).merge(table.cell(0, numberOfColumns-1))
        table.cell(1, 0).merge(table.cell(1, numberOfColumns-1))

def styleThermalTestTable(table, data = [], numberOfColumns=12, numberOfHeaderRow = 2):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    mergeCount = 0
    row = 0
    rowMergeCount = 0
    width = []
    for i in range(numberOfColumns):
        if i == 0 or i == numberOfColumns - 1:
            width.append(Cm(1.7))
        else:
            width.append(Cm(1.5))
    set_col_widths(table, columnWidth=width)
    while row < len(table.rows):
        rowHeight(table.rows[row], 0.5)
        if row < numberOfHeaderRow-1:
            table.cell(row, 0).merge(table.cell(row, numberOfColumns - 1))
            row+=1
        elif row == numberOfHeaderRow-1:
            rowHeight(table.rows[row], 0.5*4)
            cell =0
            while cell<numberOfColumns:
                if cell in data['mergeHeaderCell']:
                    table.cell(row, cell).merge(table.cell(row+1, cell))
                    cell+=1
                else:
                    table.cell(row, cell).merge(table.cell(row, cell+data['mergeHeader'][mergeCount]-1))
                    cell+=data['mergeHeader'][mergeCount]
                    mergeCount+=1
            row+=2
        elif row>numberOfHeaderRow and row<len(table.rows)-3:
            cell = 0
            while cell < numberOfColumns:
                if cell in data['tableDataMerge'][rowMergeCount][str(rowMergeCount+1)]:
                    table.cell(row, cell).merge(table.cell(row + 1, cell))
                cell+=1
            rowMergeCount+=1
            row+=2
        else:
            rowHeight(table.rows[row], 0.5 * 2)
            table.cell(row, 0).merge(table.cell(row, numberOfColumns - 1))
            row+=1

def styleDewPointTestTable(table):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.cell(0, 0).merge(table.cell(0, len(table.rows[0].cells) - 1))
    for row in table.rows:
        rowHeight(row, 0.5)

def styleCoreExcitationStateTestTable(table, numOfColumns, mergeListHorizontalHeader, mergeListVerticalHeader):
    table.cell(0, 0).merge(table.cell(0, numOfColumns - 1))
    k = 0
    for j in mergeListHorizontalHeader:
        for i in range(k, numOfColumns - 1, j):
            table.cell(1, i).merge(table.cell(1, i + j - 1))
            break
        k = k + j
    for k in mergeListVerticalHeader:
        table.cell(1, k).merge(table.cell(2, k))

def styleTable(table, data):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    mergeCount = 0
    row = 0
    dataRowCount = 1
    fieldCount = 0
    columnWidth = [Cm(1.6),Cm(2.1),Cm(1.6),Cm(1.1),Cm(1.1),Cm(1.1),Cm(1.1),Cm(1.3),Cm(2.3),Cm(1.7),Cm(2),Cm(1.2),Cm(1.6)]
    set_col_widths(table, columnWidth=columnWidth)
    while row < len(table.rows):
        rowHeight(table.rows[row], 0.5)
        if row<=data['numberOfHeaderRow']-2:
            table.cell(row, 0).merge(table.cell(row, data['numberOfColumns']-1))
            row+=1
        elif row == data['numberOfHeaderRow']-1:
            rowHeight(table.rows[row], 0.5*2)
            cell =0
            while cell < data['numberOfColumns']:
                if cell in data['mergeHeaderCell']:
                    table.cell(row, cell).merge(table.cell(row+1, cell))
                    cell+=1
                else:
                    table.cell(row, cell).merge(table.cell(row, cell+data['mergeHeader'][mergeCount]-1))
                    cell+=data['mergeHeader'][mergeCount]
                    mergeCount+=1
            row+=2
        else:
            cellCount = 0
            rowHeight(table.rows[row], 0.5*(max(data['cellTextHeight'][dataRowCount-1][str(dataRowCount)])+1))
            if str(dataRowCount) in  data['dataRowMerge'][fieldCount]:
                for cell in range(data['numberOfColumns']):
                    if cell in data['tableDataMergeCell'][fieldCount][str(dataRowCount)]:
                        table.cell(row, cell).merge(table.cell(row+data['dataRowMerge'][fieldCount][str(dataRowCount)][cellCount]-1, cell))
                        cellCount+=1
                fieldCount += 1
            dataRowCount+=1
            row+=1



def styleLeadResistance(table, numOfColumns):
    table.cell(0, 0).merge(table.cell(0, numOfColumns - 1))


def styleZeroSequence(table, numOfColumns, mergeListHorizontalHeader, mergeListVerticalHeader):

    table.cell(0, 0).merge(table.cell(0, numOfColumns - 1))

    k = 0

    for j in mergeListHorizontalHeader:

        for i in range(k, numOfColumns - 1, j):

            table.cell(1, i).merge(table.cell(1, i + j - 1))

            break

        k = k + j

    for j in mergeListVerticalHeader:
        table.cell(1, j).merge(table.cell(2, j))

