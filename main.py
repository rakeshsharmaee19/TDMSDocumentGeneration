import os
import json
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, Inches, Mm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, ns
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import test.overall
from documentCode import headers, styleTable, dataFromJSON, graph, fillTable
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from flask import Flask, request, send_from_directory
from test import impulse, potential, voltage, insulationresistance, coreexcitation, leadResistance, zeroSequence
from test import overall, bushing, firstTest, control, efficence, pdHvParallel, pdHvSeries, thermalTest, sound, dewpoint

app = Flask(__name__, static_folder='document')


@app.route("/dgaTest", methods=['GET', 'POST'])
def generateDga():
    try:
        print('API Started')
        try:
            data1 = request.data
            data = json.loads(data1)
        except Exception as e:
            return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                    "errorMessage": "DGA Data is empty"}
        document = Document()
        headerTable = headers.createHeaderFooter(document, orientation=False, )
        try:
            headerData1 = headers.headerData(data['HEADER'])
            if type(headerData1) is not list:
                return headerData1
        except Exception as e:
            return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                    "errorMessage": "Header Data is empty"}

        try:
            errorMessage = headers.fillHeaderValue(headerTable, headerData1)
            if errorMessage:
                return errorMessage
        except Exception as e:
            return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                    "errorMessage": "Header Data is not proper"}
        p = document.add_paragraph()
        styleTable.makeBoldPara(p, bold=True, italic=False, fontSize=14,
                                text='Dissolved Gas Analysis (DGA) Test Summary', alignment=1)
        p = document.add_paragraph()
        styleTable.makeBoldPara(p, bold=False, italic=False, fontSize=14,
                                text='Lab Tested : Weidman Electrical Technology', alignment=0)
        p = document.add_paragraph()
        styleTable.makeBoldPara(p, bold=False, italic=False, fontSize=14,
                                text='DGA Test Method : ASTM D-3612 Method C', alignment=0)

        numberOfRows = len(data['DATA'])
        numberOfColumns = len(data['DATA'][0])
        numberOfRowsLeft = 28
        rowValue = numberOfRows
        dataFillCount = 0
        while rowValue > 0:
            if (rowValue * 2) + 8 <= numberOfRowsLeft:
                table = document.add_table(rows=rowValue + 1, cols=numberOfColumns, style='TableGrid')
                dataDga = dataFromJSON.dgaData(data)
                styleTable.styleDgaTest(table, mergeVal=dataDga[2], widthData=dataDga[3])
                fillTable.fillDgaTable(table, dataDga, dataFillCount=dataFillCount)
                rowValue = 0
            else:
                if dataFillCount == 0:
                    table = document.add_table(rows=11, cols=numberOfColumns, style='TableGrid')
                    dataDga = dataFromJSON.dgaData(data)
                    styleTable.styleDgaTest(table, mergeVal=dataDga[2], widthData=dataDga[3])
                    counter = fillTable.fillDgaTable(table, dataDga, dataFillCount=dataFillCount)
                    dataFillCount = counter
                    rowValue -= 10
                else:
                    table = document.add_table(rows=14, cols=numberOfColumns, style='TableGrid')
                    dataDga = dataFromJSON.dgaData(data)
                    styleTable.styleDgaTest(table, mergeVal=dataDga[2], widthData=dataDga[3])
                    counter = fillTable.fillDgaTable(table, dataDga, dataFillCount=dataFillCount)
                    dataFillCount = counter
                    rowValue -= 13

        if 'FileDescription' in data.keys():
            path = fillTable.check_filepath(data)
            filename = data['FileDescription']['filename'] + '.docx'
            document.save(path + filename)
        else:
            path = 'document'
            filename = 'DGA.docx'
            document.save(path + "//" + filename)
        print('API End')
        return send_from_directory(directory=path, path=filename, as_attachment=False,
                                   mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                "errorMessage": "Data is not proper"}


@app.route("/doc", methods=['GET', 'POST'])
def genrateDoc():
    try:
        print('API Started')
        try:
            data1 = request.data
            data = json.loads(data1)
        except Exception as e:
            return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                    "errorMessage": "Heatrun Data is empty"}
        document = Document()
        headerTable = headers.createHeaderFooter(document, width=18, cell1=12, cell2=6)
        windingKey = [i for i in data['DATA'].keys()]
        gData = dataFromJSON.graphData(data, windingKey)
        try:
            headerData1 = headers.headerData(data['HEADER'])
            if type(headerData1) is not list:
                return headerData1
        except Exception as e:
            return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                    "errorMessage": "Header Data is empty"}

        try:
            errorMessage = headers.fillHeaderValue(headerTable, headerData1)
            if errorMessage:
                return errorMessage
        except Exception as e:
            return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                    "errorMessage": "Header Data is not proper"}
        winding_terminal = dataFromJSON.getTerminals(data['DATA'])
        winding = list(winding_terminal.keys())
        for i in winding_terminal:
            windingCount = 0
            otherPageGraphRow = 0
            print('Started {} Terminal'.format(i))
            topTableData = dataFromJSON.tableOneList(data['DATA'], windingCount, windingType=winding_terminal[i])
            returnDataTableData = dataFromJSON.dataTable(data['DATA'], windingCount, windingType=winding_terminal[i])
            topTable = document.add_table(rows=9, cols=len(winding_terminal[i]) + 1)
            styleTable.styleTableTop(topTable)
            if type(topTableData) is not list:
                return topTableData
            fillTable.fillTableTopData(topTable, topTableData)
            outerTable = document.add_table(rows=1, cols=2)
            styleTable.styleOuterTable(outerTable)
            innerTableData = outerTable.cell(0, 0).add_table(rows=23, cols=len(winding_terminal[i]) + 1)
            fillTable.innerTableDataFill(innerTableData, returnDataTableData[0])
            innerTableGraph = outerTable.cell(0, 1).add_table(rows=1, cols=1)
            graph.generateGraph(gData[winding_terminal[i][0]], graph.graph_heading(data, winding_terminal[i][0]))
            graph.addGraph(innerTableGraph)
            outerTableOtherPage = document.add_table(rows=1, cols=2)
            styleTable.rowHeight(outerTableOtherPage.rows[0], 15)
            styleTable.styleOuterTable(outerTableOtherPage)
            innerTableOtherPage = outerTableOtherPage.cell(0, 0).add_table(rows=22, cols=len(winding_terminal[i]) + 1)
            fillTable.innerTableDataFill(innerTableOtherPage, returnDataTableData, True)
            innerTableOtherPageGraph = outerTableOtherPage.cell(0, 1).add_table(rows=2, cols=1)
            try:
                if len(winding_terminal[i]) > 1:
                    for k in range(1, len(winding_terminal[i])):
                        graph.generateGraph(gData[winding_terminal[i][k]],
                                            graph.graph_heading(data, winding_terminal[i][k]))
                        graph.addGraph(innerTableOtherPageGraph, otherPageGraphRow)
                        otherPageGraphRow += 1
            except Exception as e:
                return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                        "errorMessage": "More than 3 terminals for {} winding".format(i)}
            if i != winding[-1]:
                document.add_paragraph('')
                document.add_paragraph('')
                document.add_paragraph('')
        try:
            print('Data Reading is completed')
            # if 'FileDescription' in data.keys():
            #     path = fillTable.check_filepath(data)
            #     filename = data['FileDescription']['filename'] + '.docx'
            #     document.save(path + filename)
            # else:
            path = 'document'
            filename = 'heatrun.docx'
            document.save(path + "//" + filename)
            # filename = data['FileDescription']['filename'] + '.docx'
        except Exception as e:
            return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                    "errorMessage": "File Description is missing"}
        print('API Ended')
        return send_from_directory(directory=path, path=filename, as_attachment=False,
                                   mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                "errorMessage": "Data is not proper"}


@app.route("/testDocument", methods=['GET', 'POST'])
def genrateTestDoc():
    try:
        data1 = request.data
        data = json.loads(data1)
        document = Document()
        section = document.sections[0]
        header = section.header
        headerTable = header.add_table(rows=6, cols=3, width=Cm(18))
        headers.headerTableStyle(headerTable)
        header.add_paragraph('')
        headerData1 = headers.headerData(data['HEADER'])
        headers.fillHeaderValue(headerTable, headerData1)
        testKey = [i for i in data['DATA'].keys()]
        numberOfRowsLeft = 35
        for i in testKey:
            if i.upper() == 'FIRSTTEST':
                print('First Test')
                try:
                    firstTest.firstTest(document, data, testKey)
                    pass
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "First or Efficience Test Data is not proper"}
            elif i.upper() == 'CONTROL':
                print('Control Test')
                try:
                    numRow = control.controlTest(document, data, numberOfRowsLeft)
                    numberOfRowsLeft = numRow
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "Control Test Data is not proper"}
            elif i.upper() == 'EFFICENCE':
                try:
                    numRow = efficence.efficenceTest(document, data, rowNumber=numberOfRowsLeft)
                    numberOfRowsLeft=numRow
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "Efficence or Regulations Test Data is not proper"}
            elif i.upper() == 'PDHVPARALLEL':
                try:
                    numRow = pdHvParallel.pdHvParallelTest(document, data=data['DATA'][i], rowNumber=numberOfRowsLeft)
                    numberOfRowsLeft=numRow
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "PDHVPARALLEL Test Data is not proper"}
            elif i.upper() == 'PDHVSERIES':
                try:
                    numRow = pdHvSeries.pdHvSeriesTest(document, data=data['DATA'][i], rowNumber=numberOfRowsLeft)
                    numberOfRowsLeft=numRow
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "PDHVSeries Test Data is not proper"}

            elif i.upper() == "INSULATINRESISTANCE":
                try:
                    numRow = insulationresistance.insulationResistanceTest(document, data=data, rowNumber=numberOfRowsLeft)
                    numberOfRowsLeft = numRow
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "Insulation Resistance Test Data is not proper"}

            elif i.upper() == 'THERMALTEST':
                try:
                    numRow = thermalTest.thermalTest(document, data=data['DATA'][i], rowNumber=numberOfRowsLeft)
                    numberOfRowsLeft=numRow
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "Thermal Test Test Data is not proper"}
            elif i.upper() == "COREEXCITATIONTABLE1":
                try:
                    numRow = coreexcitation.coreExcitation1(document, data=data, rowNumber=numberOfRowsLeft)
                    numberOfRowsLeft = numRow
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "core excitation table 1 Test Data is not proper"}
            elif i.upper() == "COREEXCITATIONTABLE2":
                try:
                    numRow = coreexcitation.coreExcitation2(document, data=data, rowNumber=numberOfRowsLeft)
                    numberOfRowsLeft = numRow
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "core excitation table 2 Test Data is not proper"}
            elif i.upper() == "ZEROSEQUENCE":
                try:
                    numRow = zeroSequence.zeroSequenceTest(document, data=data, rowNumber=numberOfRowsLeft)
                    numberOfRowsLeft = numRow
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "Zero Sequence Test Data is not proper"}

            elif i.upper() == 'SOUNDTEST':
                try:
                    numRow = sound.soundTest(document, data=data['DATA'][i], rowNumber=numberOfRowsLeft)
                    numberOfRowsLeft=numRow
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "Sound Test Test Data is not proper"}

            elif i.upper() == 'IMPULSETESTS':
                print("ImpulseTests")
                try:
                    numRow = impulse.impulseTest(document, data, numberOfRowsLeft)
                    numberOfRowsLeft = numRow
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "Impulse Test Data is not proper"}
            elif i.upper() == 'VOLTAGEPARALLEL':
                try:
                    numRow = voltage.voltageParallelTest(document, data, numberOfRowsLeft)
                    numberOfRowsLeft = numRow
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "Voltage Parallel Test Data is not proper"}
            elif i.upper() == 'VOLTAGESERIES':
                try:
                    numRow = voltage.voltageSeriesTest(document, data, numberOfRowsLeft)
                    numberOfRowsLeft = numRow
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "Voltage Series Test Data is not proper"}
            elif i.upper() == 'POTENCIALTEST':
                try:
                    numRow = potential.potencialTest(document, data, numberOfRowsLeft)
                    numberOfRowsLeft = numRow
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "Potential Test Data is not proper"}
            elif i.upper() == 'DEWPOINTTEST':
                try:
                    numRow = dewpoint.dewPointTest(document, data=data['DATA'][i], rowNumber=numberOfRowsLeft)
                    numberOfRowsLeft = numRow
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "Potential Test Data is not proper"}

            elif i.upper() == 'LEADRESISTANCE':
                try:
                    numRow = leadResistance.leadResistanceTest(document, data, rowNumber=numberOfRowsLeft)
                    numberOfRowsLeft = numRow
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "Lead Resistance Test Data is not proper"}

            elif i.upper() == 'OVERALL':
                print('overall')
                try:
                    overall.overallTest(document, data, i)
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "Overall Test Data is not proper"}
            elif i.upper() == 'BUSHING':
                print('bushing')
                try:
                    bushing.bushingTest(document, data, i)
                except Exception as e:
                    return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                            "errorMessage": "Bushing Test Data is not proper"}
            elif i.upper() == 'EXCITATION':
                print('excitation')
                countTotalTest = 0
                numberOfRowsLeft = 35
                if len(data['DATA'][i]) > 0:
                    numberOfRows = 2 + len(data['DATA'][i]['basicTempData']['basicsTemplateDataList'])
                    numberOfColumns = len(data['DATA'][i]['basicTempData']['basicsTemplateDataList'][0])
                    excitationTestData = dataFromJSON.excitationTestData(data['DATA'][i])
                    excitationTestTable = document.add_table(rows=numberOfRows, cols=numberOfColumns, style='TableGrid')
                    styleTable.styleExcitationTestTable(excitationTestTable, numberOfColumns)
                    fillTable.fillExcitationTest(excitationTestTable, excitationTestData, 0.5)
                    rowValue = numberOfRows
                    numberOfRowsLeft = numberOfRowsLeft - rowValue
                    while countTotalTest < len(data['DATA'][i]['listOfExciData']):
                        if countTotalTest == len(data['DATA'][i]['listOfExciData']) - 1 and countTotalTest == 0:
                            numberOfRows = 3 + len(
                                data['DATA'][i]['listOfExciData'][countTotalTest]['excitingDataList'])
                            numberOfColumns = len(
                                data['DATA'][i]['listOfExciData'][countTotalTest]['excitingDataList'][0])
                            excitationData = dataFromJSON.excitationData(
                                data['DATA'][i]['listOfExciData'][countTotalTest])
                            if numberOfRows + 2 <= numberOfRowsLeft:
                                fillTable.addMiddleTable(document, row=2)
                                rowValue += 2 + numberOfRows

                                merge_list = []
                                for j in range(len((data['DATA'][i]['listOfExciData'][countTotalTest][
                                    'excitingHeadersDataList'][1]))):
                                    merge_list.append(len((data['DATA'][i]['listOfExciData'][countTotalTest][
                                        'excitingHeadersDataList'][1][str(j + 1)]['2'])))
                                table = document.add_table(rows=numberOfRows, cols=numberOfColumns, style='TableGrid')
                                styleTable.styleExcitationTable(table, numberOfColumns, merge_list)
                                fillTable.fillExcitationData(table, excitationData, merge_list)

                                fillTable.addMiddleTable(document, row=(35 - rowValue))
                            else:
                                fillTable.addMiddleTable(document, row=(35 - rowValue))
                                rowValue = 0
                                numberOfRowsLeft = 35

                                merge_list = []
                                for j in range(len((data['DATA'][i]['listOfExciData'][countTotalTest][
                                    'excitingHeadersDataList'][1]))):
                                    merge_list.append(len((data['DATA'][i]['listOfExciData'][countTotalTest][
                                        'excitingHeadersDataList'][1][str(j + 1)]['2'])))
                                table = document.add_table(rows=numberOfRows, cols=numberOfColumns, style='TableGrid')
                                styleTable.styleExcitationTable(table, numberOfColumns, merge_list)
                                fillTable.fillExcitationData(table, excitationData, merge_list)
                                rowValue += numberOfRows

                                fillTable.addMiddleTable(document, row=(35 - rowValue))

                        elif countTotalTest == len(data['DATA'][i]) - 1:
                            excitationData = dataFromJSON.excitationData(
                                data['DATA'][i]['listOfExciData'][countTotalTest])
                            numberOfRows = 3 + len(
                                data['DATA'][i]['listOfExciData'][countTotalTest]['excitingDataList'])
                            numberOfColumns = len(
                                data['DATA'][i]['listOfExciData'][countTotalTest]['excitingDataList'][0])
                            if numberOfRows + 2 <= numberOfRowsLeft:
                                fillTable.addMiddleTable(document)
                                rowValue += 2 + numberOfRows

                                merge_list = []
                                for j in range(len((data['DATA'][i]['listOfExciData'][countTotalTest][
                                    'excitingHeadersDataList'][1]))):
                                    merge_list.append(len((data['DATA'][i]['listOfExciData'][countTotalTest][
                                        'excitingHeadersDataList'][1][str(j + 1)]['2'])))
                                table = document.add_table(rows=numberOfRows, cols=numberOfColumns, style='TableGrid')
                                styleTable.styleExcitationTable(table, numberOfColumns, merge_list)
                                fillTable.fillExcitationData(table, excitationData, merge_list)

                                fillTable.addMiddleTable(document, row=(35 - rowValue))

                            else:
                                fillTable.addMiddleTable(document, row=(35 - rowValue))
                                rowValue = 0
                                numberOfRowsLeft = 35

                                merge_list = []
                                for j in range(len((data['DATA'][i]['listOfExciData'][countTotalTest][
                                    'excitingHeadersDataList'][1]))):
                                    merge_list.append(len((data['DATA'][i]['listOfExciData'][countTotalTest][
                                        'excitingHeadersDataList'][1][str(j + 1)]['2'])))
                                table = document.add_table(rows=numberOfRows, cols=numberOfColumns, style='TableGrid')
                                styleTable.styleExcitationTable(table, numberOfColumns, merge_list)
                                fillTable.fillExcitationData(table, excitationData, merge_list)
                                rowValue += numberOfRows

                                fillTable.addMiddleTable(document, row=(35 - rowValue))

                        elif countTotalTest == 0:
                            excitationData = dataFromJSON.excitationData(
                                data['DATA'][i]['listOfExciData'][countTotalTest])
                            numberOfRows = 3 + len(
                                data['DATA'][i]['listOfExciData'][countTotalTest]['excitingDataList'])
                            numberOfColumns = len(
                                data['DATA'][i]['listOfExciData'][countTotalTest]['excitingDataList'][0])
                            if numberOfRows + 2 <= numberOfRowsLeft:
                                fillTable.addMiddleTable(document)
                                rowValue += 2 + numberOfRows

                                merge_list = []
                                for j in range(len((data['DATA'][i]['listOfExciData'][countTotalTest][
                                    'excitingHeadersDataList'][1]))):
                                    merge_list.append(len((data['DATA'][i]['listOfExciData'][countTotalTest][
                                        'excitingHeadersDataList'][1][str(j + 1)]['2'])))
                                table = document.add_table(rows=numberOfRows, cols=numberOfColumns, style='TableGrid')
                                styleTable.styleExcitationTable(table, numberOfColumns, merge_list)
                                fillTable.fillExcitationData(table, excitationData, merge_list)
                                if rowValue == numberOfRowsLeft:
                                    rowValue = 0
                                    numberOfRowsLeft = 35
                                else:
                                    numberOfRowsLeft = numberOfRowsLeft - numberOfRows - 2
                            else:
                                fillTable.addMiddleTable(document, row=numberOfRowsLeft)
                                rowValue = 0
                                numberOfRowsLeft = 35

                                merge_list = []
                                for j in range(len((data['DATA'][i]['listOfExciData'][countTotalTest][
                                    'excitingHeadersDataList'][1]))):
                                    merge_list.append(len((data['DATA'][i]['listOfExciData'][countTotalTest][
                                        'excitingHeadersDataList'][1][str(j + 1)]['2'])))
                                table = document.add_table(rows=numberOfRows, cols=numberOfColumns, style='TableGrid')
                                styleTable.styleExcitationTable(table, numberOfColumns, merge_list)
                                fillTable.fillExcitationData(table, excitationData, merge_list)
                                rowValue += numberOfRows

                                if rowValue == numberOfRowsLeft:
                                    rowValue = 0
                                    numberOfRowsLeft = 35
                                else:
                                    numberOfRowsLeft = numberOfRowsLeft - numberOfRows - 2
                        else:
                            numberOfRows = 3 + len(
                                data['DATA'][i]['listOfExciData'][countTotalTest]['excitingDataList'])
                            numberOfColumns = len(
                                data['DATA'][i]['listOfExciData'][countTotalTest]['excitingDataList'][0])
                            excitationData = dataFromJSON.excitationData(
                                data['DATA'][i]['listOfExciData'][countTotalTest])
                            if numberOfRows + 2 <= numberOfRowsLeft:
                                fillTable.addMiddleTable(document)

                                rowValue += 2 + numberOfRows

                                merge_list = []
                                for j in range(len((data['DATA'][i]['listOfExciData'][countTotalTest][
                                    'excitingHeadersDataList'][1]))):
                                    merge_list.append(len((data['DATA'][i]['listOfExciData'][countTotalTest][
                                        'excitingHeadersDataList'][1][str(j + 1)]['2'])))
                                table = document.add_table(rows=numberOfRows, cols=numberOfColumns, style='TableGrid')
                                styleTable.styleExcitationTable(table, numberOfColumns, merge_list)
                                fillTable.fillExcitationData(table, excitationData, merge_list)

                                if rowValue == numberOfRowsLeft:
                                    rowValue = 0
                                    numberOfRowsLeft = 35
                                else:
                                    numberOfRowsLeft = numberOfRowsLeft - numberOfRows - 2
                            else:
                                fillTable.addMiddleTable(document, row=numberOfRowsLeft)
                                rowValue = 0
                                numberOfRowsLeft = 35

                                merge_list = []
                                for j in range(len((data['DATA'][i]['listOfExciData'][countTotalTest][
                                    'excitingHeadersDataList'][1]))):
                                    merge_list.append(len((data['DATA'][i]['listOfExciData'][countTotalTest][
                                        'excitingHeadersDataList'][1][str(j + 1)]['2'])))
                                table = document.add_table(rows=numberOfRows, cols=numberOfColumns, style='TableGrid')
                                styleTable.styleExcitationTable(table, numberOfColumns, merge_list)
                                fillTable.fillExcitationData(table, excitationData, merge_list)
                                rowValue += numberOfRows

                                if rowValue == numberOfRowsLeft:
                                    rowValue = 0
                                    numberOfRowsLeft = 35
                                else:
                                    numberOfRowsLeft = numberOfRowsLeft - numberOfRows - 2
                        countTotalTest += 1
                else:
                    pass
            elif i.upper() == 'RATIO':
                print('ratio')
                countTotalTest = 0
                numberOfRowsLeft = 35
                while countTotalTest < len(data['DATA'][i]):
                    rowPass = 0
                    ratiodata = dataFromJSON.ratioTestData(data['DATA']['Ratio'][countTotalTest])
                    numberOfRows = 4 + len(data['DATA'][i][countTotalTest]['RatioData'])
                    if countTotalTest == len(data['DATA'][i]) - 1 and countTotalTest == 0:
                        numberOfRows += 2
                        heightValue = numberOfRows + 1
                        if heightValue <= numberOfRowsLeft:
                            table = document.add_table(rows=numberOfRows, cols=9, style='TableGrid')
                            styleTable.styleRatioTest(table, 9, key=True)
                            fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=True)

                            fillTable.addMiddleTable(document, row=(numberOfRowsLeft - heightValue))
                        else:
                            dataRow = numberOfRows
                            while dataRow > 0:
                                if dataRow + 6 <= 34:
                                    table = document.add_table(rows=dataRow + 6, cols=9, style='TableGrid')
                                    styleTable.styleRatioTest(table, 9, key=True)
                                    fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=True)

                                    fillTable.addMiddleTable(document, row=(numberOfRowsLeft - heightValue))
                                else:
                                    table = document.add_table(rows=34, cols=9, style='TableGrid')
                                    styleTable.styleRatioTest(table, 9, key=True)
                                    fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=True)
                                    rowPass += 28
                                    dataRow -= 28


                    elif countTotalTest == len(data['DATA'][i]) - 1:
                        if numberOfRows <= numberOfRowsLeft:
                            table = document.add_table(rows=numberOfRows, cols=9, style='TableGrid')
                            styleTable.styleRatioTest(table, 9, key=False)
                            fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=False)
                            numberOfRowsLeft -= numberOfRows

                            fillTable.addMiddleTable(document, row=(numberOfRowsLeft))
                        elif numberOfRowsLeft >= 5:
                            dataRow = numberOfRows
                            table = document.add_table(rows=numberOfRowsLeft, cols=9, style='TableGrid')
                            styleTable.styleRatioTest(table, 9, key=False)
                            fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=False)
                            rowPass = numberOfRowsLeft - 4

                            dataRow -= numberOfRowsLeft
                            numberOfRowsLeft = 35
                            while dataRow > 0:
                                if dataRow + 6 <= 34:
                                    table = document.add_table(rows=dataRow + 6, cols=9, style='TableGrid')
                                    styleTable.styleRatioTest(table, 9, key=True)
                                    fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=True)

                                    heightValue = dataRow + 7
                                    numberOfRowsLeft -= heightValue
                                    dataRow = 0
                                    fillTable.addMiddleTable(document, row=(numberOfRowsLeft))
                                else:
                                    table = document.add_table(rows=34, cols=9, style='TableGrid')
                                    styleTable.styleRatioTest(table, 9, key=True)
                                    fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=True)
                                    rowPass += 28
                                    dataRow -= 28

                        else:
                            fillTable.addMiddleTable(document, row=(numberOfRowsLeft))
                            numberOfRowsLeft = 35
                            numberOfRows += 2
                            dataRow = numberOfRows

                            while dataRow > 0:
                                if numberOfRows <= numberOfRowsLeft:
                                    table = document.add_table(rows=numberOfRows, cols=9, style='TableGrid')
                                    styleTable.styleRatioTest(table, 9, key=True)
                                    fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=True)
                                    heightValue = dataRow + 1
                                    numberOfRowsLeft -= heightValue
                                    dataRow = 0
                                else:
                                    if dataRow + 6 <= 34:
                                        table = document.add_table(rows=dataRow + 6, cols=9, style='TableGrid')
                                        styleTable.styleRatioTest(table, 9, key=True)
                                        fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=True)

                                        heightValue = dataRow + 7
                                        numberOfRowsLeft -= heightValue
                                        dataRow = 0
                                    else:
                                        table = document.add_table(rows=34, cols=9, style='TableGrid')
                                        styleTable.styleRatioTest(table, 9, key=True)
                                        fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=True)
                                        rowPass += 28
                                        dataRow -= 28
                                fillTable.addMiddleTable(document, row=(numberOfRowsLeft))

                    elif countTotalTest == 0:
                        numberOfRows += 2
                        heightValue = numberOfRows + 1
                        if heightValue <= numberOfRowsLeft:
                            table = document.add_table(rows=numberOfRows, cols=9, style='TableGrid')
                            styleTable.styleRatioTest(table, 9, key=True)
                            fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=True)
                        else:
                            dataRow = numberOfRows
                            while dataRow > 0:
                                if dataRow + 6 <= 34:
                                    table = document.add_table(rows=dataRow + 6, cols=9, style='TableGrid')
                                    styleTable.styleRatioTest(table, 9, key=True)
                                    fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=True)

                                    heightValue = dataRow + 7
                                    numberOfRowsLeft -= heightValue
                                    dataRow = 0
                                else:
                                    table = document.add_table(rows=34, cols=9, style='TableGrid')
                                    styleTable.styleRatioTest(table, 9, key=True)
                                    fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=True)
                                    rowPass += 28
                                    if dataRow == numberOfRows:
                                        dataRow -= 34
                                    else:
                                        dataRow -= 28
                    else:
                        if numberOfRows <= numberOfRowsLeft:
                            table = document.add_table(rows=numberOfRows, cols=9, style='TableGrid')
                            styleTable.styleRatioTest(table, 9, key=False)
                            fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=False)
                            numberOfRowsLeft -= numberOfRows

                        elif numberOfRowsLeft >= 5:
                            dataRow = numberOfRows
                            table = document.add_table(rows=numberOfRowsLeft, cols=9, style='TableGrid')
                            styleTable.styleRatioTest(table, 9, key=False)
                            fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=False)
                            rowPass = numberOfRowsLeft - 4
                            dataRow -= numberOfRowsLeft
                            numberOfRowsLeft = 35
                            while dataRow > 0:
                                if dataRow + 6 <= 34:
                                    table = document.add_table(rows=dataRow + 6, cols=9, style='TableGrid')
                                    styleTable.styleRatioTest(table, 9, key=True)
                                    fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=True)

                                    heightValue = dataRow + 7
                                    numberOfRowsLeft -= heightValue
                                    dataRow = 0
                                else:
                                    table = document.add_table(rows=34, cols=9, style='TableGrid')
                                    styleTable.styleRatioTest(table, 9, key=True)
                                    fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=True)
                                    rowPass += 28
                                    dataRow -= 28

                        else:
                            fillTable.addMiddleTable(document, row=(numberOfRowsLeft))

                            numberOfRowsLeft = 35
                            numberOfRows += 2
                            dataRow = numberOfRows
                            while dataRow > 0:
                                if dataRow + 6 <= 34:
                                    table = document.add_table(rows=dataRow + 6, cols=9, style='TableGrid')
                                    styleTable.styleRatioTest(table, 9, key=True)
                                    fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=True)

                                    heightValue = dataRow + 7
                                    numberOfRowsLeft -= heightValue
                                    dataRow = 0
                                else:
                                    table = document.add_table(rows=34, cols=9, style='TableGrid')
                                    styleTable.styleRatioTest(table, 9, key=True)
                                    fillTable.fillRatioTest(table, ratiodata, rowValue=rowPass, rowH=0.5, key=True)
                                    rowPass += 28

                                    if dataRow == numberOfRows:
                                        dataRow -= 34
                                    else:
                                        dataRow -= 28
                    countTotalTest += 1
        # path1 = f'C:\\Users\\Rahul G\\PycharmProjects\\TDMSDocumentGeneration\\document\\'
        # dir_list = os.listdir(path1)
        # if len(dir_list) > 0:
        #     os.remove(path1 + "/" + dir_list[0])
        path1 = 'document'
        # dir_list = os.listdir(path1)
        # if len(dir_list) > 0:
        #     os.remove(path1 + "//" + dir_list[0])
        filename = 'testDocument.docx'
        document.save(path1 + "//" + filename)
        return send_from_directory(directory='document', path=filename, as_attachment=False,
                                   mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return {"status": "INTERNAL_SERVER_ERROR", "statusCode": 500, "message": "failed", "error": str(e),
                "errorMessage": "Data is not proper"}


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=1048, debug=True)
