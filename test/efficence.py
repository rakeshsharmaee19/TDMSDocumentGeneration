from docx.shared import Cm
from documentCode import fillTable, dataFromJSON, styleTable
from docx.enum.table import  WD_CELL_VERTICAL_ALIGNMENT

def efficenceTest(document, data, rowNumber= 0):
    print('Efficence Test')
    countTotalTest = 0
    numberOfRowsLeft = rowNumber
    while countTotalTest < len(data['DATA']['efficence']):
        numberOfRows = len(data['DATA']['efficence'][countTotalTest]['headers'])+len(data['DATA']['efficence'][countTotalTest]['data'])
        numberOfColumns = len(data['DATA']['efficence'][countTotalTest]['data'][0])
        if numberOfRowsLeft==35:
            efficenceOuterTable = document.add_table(rows=1, cols=2, style='TableGrid')
            styleTable.styleEfficenceOuterTable(efficenceOuterTable, row=numberOfRows)
            efficenceTbaleData = dataFromJSON.efficenceTestData(data['DATA']['efficence'][countTotalTest])
            efficenceTable = efficenceOuterTable.cell(0, 0).add_table(rows=numberOfRows, cols=numberOfColumns)
            styleTable.styleEfficenceTable(table=efficenceTable)
            efficenceOuterTable.cell(0, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            fillTable.fillEfficenceTestTable(efficenceTable, efficenceTbaleData)
            regulationTableData = dataFromJSON.efficenceTestData(data['DATA']['regulations'][countTotalTest])
            regulationTable = efficenceOuterTable.cell(0, 1).add_table(rows=numberOfRows, cols=numberOfColumns)
            styleTable.styleEfficenceTable(table=regulationTable)
            efficenceOuterTable.cell(0, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            fillTable.fillEfficenceTestTable(regulationTable, regulationTableData)
        else:
            if numberOfRows+2<=numberOfRowsLeft:
                efficenceOuterTable = document.add_table(rows=1, cols=2)
                styleTable.styleEfficenceOuterTable(efficenceOuterTable, row=numberOfRows)
                efficenceTbaleData = dataFromJSON.efficenceTestData(data['DATA']['efficence'][countTotalTest])
                efficenceTable = efficenceOuterTable.cell(0,0).add_table(rows=numberOfRows, cols=numberOfColumns)
                styleTable.styleEfficenceTable(table=efficenceTable)
                efficenceOuterTable.cell(0, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                fillTable.fillEfficenceTestTable(efficenceTable, efficenceTbaleData)
                regulationTableData = dataFromJSON.efficenceTestData(data['DATA']['regulations'][countTotalTest])
                regulationTable = efficenceOuterTable.cell(0,1).add_table(rows=numberOfRows, cols=numberOfColumns)
                styleTable.styleEfficenceTable(table=regulationTable)
                efficenceOuterTable.cell(0, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                fillTable.fillEfficenceTestTable(regulationTable, regulationTableData)
            else:
                fillTable.styleEfficenceTable(document, row=numberOfRowsLeft)
                numberOfRowsLeft=35
                efficenceOuterTable = document.add_table(rows=1, cols=2)
                styleTable.styleEfficenceOuterTable(efficenceOuterTable, row=numberOfRows)
                efficenceTbaleData = dataFromJSON.efficenceTestData(data['DATA']['efficence'][countTotalTest])
                efficenceTable = efficenceOuterTable.cell(0, 0).add_table(rows=numberOfRows, cols=numberOfColumns)
                styleTable.styleEfficenceTable(table=efficenceTable)
                efficenceOuterTable.cell(0, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                fillTable.fillEfficenceTestTable(efficenceTable, efficenceTbaleData)
                regulationTableData = dataFromJSON.efficenceTestData(data['DATA']['regulations'][countTotalTest])
                regulationTable = efficenceOuterTable.cell(0, 1).add_table(rows=numberOfRows, cols=numberOfColumns)
                styleTable.styleEfficenceTable(table=regulationTable)
                efficenceOuterTable.cell(0, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                fillTable.fillEfficenceTestTable(regulationTable, regulationTableData)
        numberOfRowsLeft-=numberOfRows+2
        countTotalTest+=1
    return  numberOfRowsLeft
