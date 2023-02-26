from docx.shared import Cm

from documentCode import fillTable, dataFromJSON, styleTable

def firstTest(document, data, testKey):
    numberOfRowsLeft = 35
    rowValue = 0
    for i in testKey:
        if i.upper() == 'FIRSTTEST':
            print('First Test')
            numbersOfRows = 1+len(data['DATA'][i][0]['Data'])
            numberSOfColumns = len(data['DATA'][i][0]['Data'][0])
            firstTable = document.add_table(rows=numbersOfRows, cols=numberSOfColumns, style='TableGrid')
            firstTableData = dataFromJSON.firstTestData(data['DATA'][i][0])
            styleTable.set_col_widths(firstTable, columnWidth=firstTableData[2])
            fillTable.fillFirstTestData(firstTable, firstTableData)
            rowValue += numbersOfRows
            numberOfRowsLeft-= rowValue

        elif i.upper() == 'EFFICIENCE':
            print('Efficience Test')
            efficienceData = dataFromJSON.EfficeinceTestData(data['DATA'][i])
            numbersOfRowsTableOne = 1+len(efficienceData[1])
            numbersOfRowsTableTwo = 2+ len(data['DATA'][i][0]["Table"]['Data'])
            numbersOfColumnsTwo = len(data['DATA'][i][0]["Table"]['Data'][0])
            textRow = sum(efficienceData[2])
            if (textRow+6+len(data['DATA'][i][0]["Table"]['Data'])*2)<=numberOfRowsLeft:
                # section = document.sections[0]
                # section.left_margin = Cm(2)
                fillTable.addMiddleTable(document, row=1)
                textTable = document.add_table(rows=numbersOfRowsTableOne, cols=1)
                styleTable.styleTextTable(document, textTable, efficienceData[2])
                fillTable.filltextTable(textTable, efficienceData[0]+efficienceData[1])
                fillTable.addMiddleTable(document, row=1)

                rowValue = textRow+6+len(data['DATA'][i][0]["Table"]['Data'])*2
                mergerRow = dataFromJSON.merge_list(data)
                dataTable = document.add_table(rows=numbersOfRowsTableTwo, cols=numbersOfColumnsTwo, style='TableGrid')
                styleTable.styleEfficienceTable(dataTable, numberOfRows=numbersOfRowsTableTwo, numberOfColumns=numbersOfColumnsTwo, merge_list=mergerRow)

                fillTable.fillEfficienceTestTable(dataTable, data=[efficienceData[3], efficienceData[4]], mergerValue = mergerRow)
                fillTable.addMiddleTable(document, row=numberOfRowsLeft-rowValue)
        else:
            pass

