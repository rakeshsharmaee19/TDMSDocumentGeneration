from docx.shared import Cm
from documentCode import fillTable, dataFromJSON, styleTable

def soundTest(document, data={}, rowNumber = 0):
    print('Sound Test')
    numberOfRowsLeft = rowNumber
    countTotalTest = 0
    while countTotalTest < len(data):
        soundTestData = dataFromJSON.soundTestData(data[0])
        numberOfRows = soundTestData['numberOfHeaderRow']+soundTestData['numberOfRows']+1
        numberOfColumns = soundTestData['numberOfColumns']
        value = 0
        for j in range(soundTestData['numberOfRows']):
            value = value+(max(soundTestData['cellTextHeight'][j][str(j+1)]))
        rowValue = soundTestData['numberOfHeaderRow']+ value+1
        if numberOfRowsLeft == 35:
            soundTestTable = document.add_table(rows = numberOfRows, cols=numberOfColumns, style='TableGrid')
            styleTable.styleTable(soundTestTable, soundTestData)

        else:
            if rowValue+2<=numberOfRowsLeft:
                soundTestTable = document.add_table(rows=numberOfRows, cols=numberOfColumns, style='TableGrid')
                styleTable.styleTable(soundTestTable, soundTestData)
                fillTable.fillTable(soundTestTable, soundTestData, numberOfHeaderRows=soundTestData['numberOfHeaderRow'], numberOfCells=soundTestData['numberOfColumns'])
            else:
                fillTable.addMiddleTable(document, row=numberOfRowsLeft)
                numberOfRowsLeft=35
                soundTestTable = document.add_table(rows=numberOfRows, cols=numberOfColumns, style='TableGrid')
                styleTable.styleTable(soundTestTable, soundTestData)
                fillTable.fillTable(soundTestTable, soundTestData, numberOfHeaderRows=soundTestData['numberOfHeaderRow'],
                                    numberOfCells=soundTestData['numberOfColumns'])

        numberOfRowsLeft-=(rowValue+1+soundTestData['numberOfRows'])
        countTotalTest+=1
    return numberOfRowsLeft